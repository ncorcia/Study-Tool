"""Text extraction from DOCX files using python-docx."""

import docx


def extract_text(file_path: str) -> str:
    """Extract text from a DOCX file, including paragraph and table content."""
    document = docx.Document(file_path)

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n\n".join(parts).strip()
